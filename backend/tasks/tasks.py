from django.core.files.base import ContentFile
from matplotlib.ticker import ScalarFormatter

import uuid
from celery import shared_task
from celery.utils.log import get_task_logger
logger = get_task_logger(__name__)
from dateutil.parser import parse
import os
from django.core.management.base import BaseCommand
from django.db import connections
from datetime import datetime, timedelta
from azure.storage.blob import BlobServiceClient, BlobClient, ContainerClient
import subprocess
from notifications.models import Notifications, CalendarEvents
from notifications.serializers import NotificationsSerializer
from logs.models import Log, LogKPIs
from logs.serializers import LogSerializer, LogKPIsSerializer, LogContentSerializer 
from data.models import UserAccount, user_profile, region_manager, DisclosuresProductProviders, DisclosuresAdvisorSubCodes
from data.serializers import DisclosuresProductProviders_Serializer, DisclosuresAdvisorSubCodes_Serializer
from django.db.models import Q, Sum, Count, FloatField
from django.db.models.functions import Cast
from compliance.models import ComplianceDocument, DocumentComments, arc, GateKeeping
import base64
import pandas as pd
from django.conf import settings
import environ
from docx.shared import Pt, RGBColor, Inches
from docx import Document
import matplotlib.pyplot as plt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
from io import BytesIO
import os
from data.models import regions, UserAccount, Disclosures, user_profile, region_manager
import babel.numbers
env = environ.Env(
    # set casting, default value
    DEBUG=(bool, False)
)
from django.conf import settings
environ.Env.read_env(os.path.join(settings.BASE_DIR, '.env'))


def currency_formatter(x, pos):
    return f'R{x:,.0f}'

class DB_Backup_Command(BaseCommand):
    help = 'Backup PostgreSQL database and upload to Azure Blob'

    def handle(self, *args, **options):
        admin_ids = []
        users = UserAccount.objects.filter(is_superuser=True)
        db_name = connections['default'].settings_dict['NAME']
        db_user = connections['default'].settings_dict['USER']
        os.environ['PGPASSWORD'] = connections['default'].settings_dict['PASSWORD']
        db_password = connections['default'].settings_dict['PASSWORD']
        backup_dir = ''
        timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
        backup_file = f'{backup_dir}/{db_name}_backup_{timestamp}_{uuid.uuid4()}.sql'
        log_user_data = []
        for user in users:
            title = f"Database Backup Status"
            notificationData = {
                "account" : user.pk,
                "notificationType" : 4,
                
                "user" : user.pk,
                "title" : title,
                "message" : f"Database backup and uploading.",
                "status" : False,
            }
            serializer = NotificationsSerializer(data=notificationData)
            if serializer.is_valid():
                serializer.create(serializer.validated_data)
            else:
                logger.info(serializer.errors)
            logData = {
                "account" : user.pk,
                "log_type": 5,
                "log_name" : f"Database backup task {uuid.uuid4()}.",
            }
            logSerializer = LogSerializer(data=logData)
            if logSerializer.is_valid():
                log = logSerializer.create(logSerializer.validated_data)
                logger.info(f"Log created for database backup task {log.pk}")
                logId = log.pk
                log_user_data.append({
                    "log_id" : logId,
                    "user": user.pk
                })
        # Perform backup using pg_dump
        subprocess.run(['pg_dump', '-h', 'localhost', '-U', db_user, '-d', db_name, '-F', 'c', '-f', backup_file])
        del os.environ['PGPASSWORD']

        self.stdout.write(self.style.SUCCESS(f'Successfully created backup: {backup_file}'))

        for user in log_user_data:    
            logContent = {
                "account" : user['user'],
                "log" : user['log_id'],
                "log_type" : 1,
                "log_description" : f"Database backup created at {backup_file}",
            }
            logContentSerializer = LogContentSerializer(data=logContent)
            if logContentSerializer.is_valid():
                logContentSerializer.create(logContentSerializer.validated_data)
            else:
                logger.info(logContentSerializer.errors)


        # Upload to Azure Blob
        blob_service_client = BlobServiceClient(account_url=f"https://{settings.AZURE_ACCOUNT_NAME}.blob.core.windows.net/", credential=settings.AZURE_ACCOUNT_KEY)
        container_client = blob_service_client.get_container_client(settings.AZURE_CONTAINER)
        blob_client = container_client.get_blob_client(f'backups/{os.path.basename(backup_file)}')

        for user in log_user_data:
            logContent = {
                "account" : user['user'],
                "log" : user['log_id'],
                "log_type" : 2,
                "log_description" : f"Connection made to backup server.",
            }
            logContentSerializer = LogContentSerializer(data=logContent)
            if logContentSerializer.is_valid():
                logContentSerializer.create(logContentSerializer.validated_data)
            else:
                logger.info(logContentSerializer.errors)
            logContent = {
                "account" : user['user'],
                "log" : user['log_id'],
                "log_type" : 3,
                "log_description" : f"Uploading the backup {backup_file} to backup server.",
            }
            logContentSerializer = LogContentSerializer(data=logContent)
            if logContentSerializer.is_valid():
                logContentSerializer.create(logContentSerializer.validated_data)
            else:
                logger.info(logContentSerializer.errors)
        try:
            with open(backup_file, "rb") as data:
                blob_client.upload_blob(data)
            for user in log_user_data:
                logContent = {
                    "account" : user['user'],
                    "log" : user['log_id'],
                    "log_type" : 4,
                    "log_description" : f"Backup {backup_file} uploaded to backup server.",
                }
                logContentSerializer = LogContentSerializer(data=logContent)
                if logContentSerializer.is_valid():
                    logContentSerializer.create(logContentSerializer.validated_data)
                else:
                    logger.info(logContentSerializer.errors)

            os.remove(backup_file)

            self.stdout.write(self.style.SUCCESS('Backup uploaded to Azure Blob'))
            for user in log_user_data:
                notificationData = {
                    "account" : user['user'],
                    "notificationType" : 4,
                    
                    
                    "log": "",
                    "user" : user['user'],
                    "title" : title,
                    "message" : f"Database backup and uploading to server completed.",
                    "status" : False,
                }
                serializer = NotificationsSerializer(data=notificationData)
                if serializer.is_valid():
                    serializer.create(serializer.validated_data)
                else:
                    logger.info(serializer.errors)
        except:
            for user in log_user_data:
                logContent = {
                    "account" : user['user'],
                    "log" : user['log_id'],
                    "log_type" : 4,
                    "log_description" : f"Backup {backup_file} uploading to backup server failed.",
                }
                logContentSerializer = LogContentSerializer(data=logContent)
                if logContentSerializer.is_valid():
                    logContentSerializer.create(logContentSerializer.validated_data)
                else:
                    logger.info(logContentSerializer.errors)
            logger.info()


        blob_client.close()

        for user in log_user_data:
            logContent = {
                "account" : user['user'],
                "log" : user['log_id'],
                "log_type" : 6,
                "log_description" : f"Connection closed to backup server.",
            }
            logContentSerializer = LogContentSerializer(data=logContent)
            if logContentSerializer.is_valid():
                logContentSerializer.create(logContentSerializer.validated_data)
            else:
                logger.info(logContentSerializer.errors)

            title = f"Database Backup Status"
            notificationData = {
                "account" : user['user'],
                "notificationType" : 4,
                
                
                "log": user['log_id'],
                "user" : user['log_id'],
                "title" : title,
                "message" : f"Database backup and uploading to your backup server completed.",
                "status" : False,
            }
            serializer = NotificationsSerializer(data=notificationData)
            if serializer.is_valid():
                serializer.create(serializer.validated_data)
            else:
                logger.info(serializer.errors)
            log = Log.objects.get(id=user['log_id'], account=user['user'])
            logData = {
                'status' : 1,
                "closed_at": datetime.now()
            }
            logSerializer = LogSerializer(instance=log,data=logData, partial=True)
            if logSerializer.is_valid():
                logSerializer.save()
            else:
                logger.info(logSerializer.errors)

class Delete_DB_Backup(BaseCommand):
    help = 'Delete backups older than 15 days from Azure Blob'

    def handle(self, *args, **options):
        
        max_age_days = 30

        # Calculate the date threshold (one month ago)
        one_month_ago = datetime.utcnow() - timedelta(days=max_age_days)

        # Connect to Azure Blob
        blob_service_client = BlobServiceClient(account_url=f"https://{settings.AZURE_ACCOUNT_NAME}.blob.core.windows.net/", credential=settings.AZURE_ACCOUNT_KEY)
        container_client = blob_service_client.get_container_client(settings.AZURE_CONTAINER)
    

        # List objects in the Blob Container
        blob_list = container_client.list_blobs('backups/')
        
        # Iterate over the objects and delete those older than the threshold
        for blob in blob_list:
            if blob.last_modified < one_month_ago:
                # Delete the blob
                container_client.delete_blob(blob.name)
                self.stdout.write(self.style.SUCCESS(f'Deleted backup from Blob: {blob.name}'))
        

from django.core.management import call_command
@shared_task
def upload_backup():
    call_command(DB_Backup_Command())

@shared_task
def delete_backup():

    call_command(Delete_DB_Backup())

@shared_task
def events():

    events = CalendarEvents.objects.filter(start_time=(datetime.today() + timedelta(days=1)).date())
    for event in events:
        if "cut-off" in str(event.title).lower():
            users = UserAccount.objects.all()
            event_name = ""
            if "afp" in str(event.title).lower():
                users = users.filter(Q(email__icontains="sanlam")|Q(email__icontains="kaspar")|Q(email__icontains="yahoo")|Q(email__icontains="gmail"))
                event_name = "AFP"
            if 'sfp' in str(event.title).lower():
                users = users.filter(Q(email__icontains="sfp")|Q(email__icontains="kaspar")|Q(email__icontains="yahoo")|Q(email__icontains="gmail"))
                event_name = "SFP"
            for user in users:
                user_name = user.first_name + " " + user.last_name
                user_profile_data = user_profile.objects.filter(user=user.pk)
                if user_profile_data.exists():
                    user_profile_data = user_profile_data.first()
                    user_name = user_profile_data.Full_Name
                data = {
                    "title" : event.title,
                    "notificationType" : 2,
                    "account" : user.pk,
                    "message" : f"""<div style="max-width: 600px; margin: 20px auto; padding: 20px; background-color: #fff; border-radius: 8px; box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);">
                        <h1 style="color: #333; margin-bottom: 20px;">📢 Important Announcement: {event.title}</h1>
                        <div style="background-color: #6AC7D2; color: #fff; padding: 15px; border-radius: 4px; margin-bottom: 20px;">
                            <p style="margin: 0;">Dear {user_name},</p>
                            <p style="margin: 0;">We would like to bring to your attention an important event: tomorrow marks the {event.title} Date. This date is a significant milestone in our project timeline and requires everyone's attention and cooperation to ensure its smooth execution.</p>
                        </div>
                        <p style="color: #666; line-height: 1.6; margin-bottom: 15px;"><strong>Event Details:</strong></p>
                        <ul style="color: #666; line-height: 1.6; margin-bottom: 15px;">
                            <li><strong>Title:</strong> {event.title}</li>
                            <li><strong>Date:</strong> {event.start_time.strftime('%d %b %Y')}</li>
                            <li><strong>End Date:</strong> {event.end_time.strftime('%d %b %Y')}</li>
                        </ul>
                        <p style="color: #666; line-height: 1.6; margin-bottom: 15px;"><strong>Actions Required:</strong></p>
                        <ol style="color: #666; line-height: 1.6; margin-bottom: 15px;">
                            <li>Please be reminded of the {event_name} cut-off tomorrow. You got this!.</li>
                        </ol>
                        <p style="color: #666; line-height: 1.6; margin-bottom: 15px;"><strong>Need Assistance?</strong></p>
                        <p style="color: #666; line-height: 1.6; margin-bottom: 15px;">If you have any questions or require assistance regarding the {event.title} or related tasks, please don't hesitate to reach out to your project manager or team lead for guidance.</p>
                        <p style="color: #666; line-height: 1.6; margin-bottom: 15px;">Thank you for your attention to this matter, and let's work together to ensure a successful {event.title}!</p>
                        <p style="color: #666; line-height: 1.6; margin-bottom: 15px;">Best regards,<br>Succession</p>

                    </div>"""
                }
                notificationSerializer = NotificationsSerializer(data=data)
                if notificationSerializer.is_valid():
                    notificationSerializer.create(notificationSerializer.validated_data)
                else:
                    logger.info(notificationSerializer.errors)
        if "pay run" in str(event.title).lower() or "pay-run" in str(event.title).lower():
            users = UserAccount.objects.all()
            event_name = ""
            if "afp" in str(event.title).lower():
                users = users.filter(Q(email__icontains="sanlam")|Q(email__icontains="kaspar")|Q(email__icontains="yahoo")|Q(email__icontains="gmail"))
                event_name = "AFP"
            if 'sfp' in str(event.title).lower():
                users = users.filter(Q(email__icontains="sfp")|Q(email__icontains="kaspar")|Q(email__icontains="yahoo")|Q(email__icontains="gmail"))
                event_name = "SFP"
    
            for user in users:
                user_name = user.first_name + " " + user.last_name
                user_profile_data = user_profile.objects.filter(user=user.pk)
                if user_profile_data.exists():
                    user_profile_data = user_profile_data.first()
                    user_name = user_profile_data.Full_Name
                data = {
                    "title" : event.title,
                    "notificationType" : 2,
                    "account" : user.pk,
                    "message" : f"""
                        <div style="max-width: 600px; margin: 20px auto; padding: 20px; background-color: #fff; border-radius: 8px; box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);">
                            <h1 style="color: #333; margin-bottom: 20px;">💰 Pay Run Announcement 💵</h1>
                            <div style="background-color: #FFCC00; color: #333; padding: 15px; border-radius: 4px; margin-bottom: 20px;">
                                <p style="margin: 0;">Dear {user_name},</p>
                                <p style="margin: 0;">This is a friendly reminder that the next pay run 🤑 is scheduled for tomorrow {event.start_time.strftime('%d %b %Y')}.</p>
                            </div>
                            <p style="color: #666; line-height: 1.6; margin-bottom: 15px;">If you have any questions or concerns regarding the pay run, feel free to reach out.</p>
                            <p style="color: #666; line-height: 1.6; margin-bottom: 15px;">Best regards,<br>Succession</p>
                        </div>"""
                }
                notificationSerializer = NotificationsSerializer(data=data)
                if notificationSerializer.is_valid():
                    notificationSerializer.create(notificationSerializer.validated_data)
                else:
                    logger.info(notificationSerializer.errors)
        if "query" in str(event.title).lower():
            users = UserAccount.objects.all()
            event_name = ""
            if "afp" in str(event.title).lower():
                users = users.filter(Q(email__icontains="sanlam")|Q(email__icontains="kaspar")|Q(email__icontains="yahoo")|Q(email__icontains="gmail"))
                event_name = "AFP"
            if 'sfp' in str(event.title).lower():
                users = users.filter(Q(email__icontains="sfp")|Q(email__icontains="kaspar")|Q(email__icontains="yahoo")|Q(email__icontains="gmail"))
                event_name = "SFP"
    
            for user in users:
                user_name = user.first_name + " " + user.last_name
                user_profile_data = user_profile.objects.filter(user=user.pk)
                if user_profile_data.exists():
                    user_profile_data = user_profile_data.first()
                    user_name = user_profile_data.Full_Name
                data = {
                    "title" : event.title,
                    "notificationType" : 2,
                    "account" : user.pk,
                    "message" : f"""
                        <div style="max-width: 600px; margin: 20px auto; padding: 20px; background-color: #fff; border-radius: 8px; box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);">
                            <h1 style="color: #333; margin-bottom: 20px;">Query Day Annoucement</h1>
                            <div style="background-color: #FFCC00; color: #333; padding: 15px; border-radius: 4px; margin-bottom: 20px;">
                                <p style="margin: 0;">Dear {user_name},</p>
                                <p style="margin: 0;">This is a friendly reminder that the next query is scheduled for tomorrow {event.start_time.strftime('%d %b %Y')}. Please submit any pay run related query by tomorrow to ensure it's resolved before the pay run finalisation.</p>
                            </div>
                            <p style="color: #666; line-height: 1.6; margin-bottom: 15px;">If you have any questions or concerns regarding the pay run, feel free to reach out.</p>
                            <p style="color: #666; line-height: 1.6; margin-bottom: 15px;">Best regards,<br>Succession</p>
                        </div>"""
                }
                notificationSerializer = NotificationsSerializer(data=data)
                if notificationSerializer.is_valid():
                    notificationSerializer.create(notificationSerializer.validated_data)
                else:
                    logger.info(notificationSerializer.errors)
        if "birthday" in str(event.title).lower():
            users = UserAccount.objects.all()
            for user in users:
                user_name = user.first_name + " " + user.last_name
                user_profile_data = user_profile.objects.filter(user=user.pk)
                if user_profile_data.exists():
                    user_profile_data = user_profile_data.first()
                    user_name = user_profile_data.Full_Name
                data = {
                    "title" : event.title,
                    "notificationType" : 2,
                    "account" : user.pk,
                    "message" : f"""<div style="max-width: 600px; margin: 20px auto; padding: 20px; background-color: #fff; border-radius: 8px; box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);">
                        <h1 style="color: #333; margin-bottom: 20px;">Birthday ({event.title}) Celebration Announcement 🎂</h1>
                        <div style="background-color: #6AC7D2; color: #fff; padding: 15px; border-radius: 4px; margin-bottom: 20px;">
                            <p style="margin: 0;">Dear {user_name},</p>
                            <p style="margin: 0;">Please join us in wishing a very happy birthday to {event.title}! 🎈🎂 Let's make their day special with your warm wishes and greetings!</p>
                        </div>
                        <p style="color: #666; line-height: 1.6; margin-bottom: 15px;">Best regards,<br>Succession</p>

                        
                    </div>"""
                }
                notificationSerializer = NotificationsSerializer(data=data)
                if notificationSerializer.is_valid():
                    notificationSerializer.create(notificationSerializer.validated_data)
                else:
                    logger.info(notificationSerializer.errors)
            
    return True



@shared_task
def compliance_export(userId, data):
    user = UserAccount.objects.get(id=userId)
    title = "Advise Monitoring Export"
    message = "Advise Monitoring complete data export"
    documents = ComplianceDocument.objects.all().order_by('-created_at')
    filter_type = int(data['filter_type'])
    file_initial_name = "all"
    if filter_type != 0:
        if filter_type == 1:
            year = data['year']
            file_initial_name = year
            message = f"Advise Monitoring complete data export for year {year}"
            documents = documents.filter(created_at__year=year).order_by('created_at__year')
        elif filter_type == 2:
            date = str(data['month_date'])
            month = date.split('-')[1]
            year = date.split('-')[0]
            file_initial_name = year + "_" + month
            range = parse(f"{year}-{month}").strftime("%B %Y")
            message = f"Advise Monitoring complete data export for {range}"
            documents = documents.filter(created_at__year=year, created_at__month=month).order_by('created_at__year', 'created_at__month')
        elif filter_type == 3:
            date = str(data['date'])
            day = date.split('-')[2]
            month = date.split('-')[1]
            year = date.split('-')[0]
            file_initial_name = year + "_" + month + "_" + day
            range = parse(f"{year}-{month}-{day}").strftime("%d %B %Y")
            message = f"Advise Monitoring complete data export for {range}"
            documents = documents.filter(created_at__year=year, created_at__month=month, created_at__day=day).order_by('created_at__year', 'created_at__month', 'created_at__day')
        else:
            fromdate = data['fromdate']
            todate = data['todate']
            file_initial_name = fromdate + "to" + todate
            starting_date = parse(f"{fromdate}").strftime("%d %B %Y")
            ending_date = parse(f"{todate}").strftime("%d %B %Y")
            message = f"Advise Monitoring complete data export of range which starts from {starting_date} and ends on {ending_date}"
            date_range = (datetime.strptime(fromdate, '%Y-%m-%d') , datetime.strptime(todate, '%Y-%m-%d') + timedelta(days=1))
            documents = documents.filter(created_at__range=date_range)
    requestedData = []
    notificationData = {
        "account" : user.pk,
        "notificationType" : 6,        
        "user" : user.pk,
        "title" : title,
        "message" : message,
        "status" : False,
    }
    serializer = NotificationsSerializer(data=notificationData)
    if serializer.is_valid():
        serializer.create(serializer.validated_data)
    else:
        logger.info(serializer.errors)
    for document in documents:
        advisor_name = document.advisor.first_name + " " + document.advisor.last_name
        advisor_id = document.IdNumber
        region_name = "N/A"
        regional_manager_name = "N/A"
        regional_manager_email = "N/A"
        bac_name = ""
        bac_email = ""
        advisor = user_profile.objects.filter(user=document.advisor.pk)
        if advisor.exists():
            advisor = advisor.first()
            advisor_name = advisor.Full_Name 
            advisor_id = advisor.ID_Number 
            if advisor.region:
                region_name = advisor.region.region
                regional_manager = region_manager.objects.filter(region=advisor.region.pk)
                if regional_manager.exists():
                    regional_manager = regional_manager.first()
                    regional_manager = UserAccount.objects.filter(id=regional_manager.manager.pk)
                    if regional_manager.exists():
                        regional_manager = regional_manager.first()
                        regional_manager_name = regional_manager.first_name + " " + regional_manager.last_name
                        regional_manager_email = regional_manager.email
                        regional_manager_profile = user_profile.objects.filter(user=regional_manager.pk)
                        if regional_manager_profile.exists():
                            regional_manager_name = regional_manager_profile.first().Full_Name
            bac = advisor.bac.pk if advisor.bac else ""
            if bac != "":
                bac = UserAccount.objects.get(id=bac)
                bac_name = bac.first_name + " " + bac.last_name
                bac_email = bac.email
                bac_details = user_profile.objects.filter(id=bac.pk)
                if bac_details.exists():
                    bac_details = bac_details.first()
                    bac_name = bac_details.Full_Name
        business_type = document.businessType
        business_type_name = ""
        if business_type == 1:
            business_type_name = "Business Assurance"
        if business_type == 2:
            business_type_name = "Comm release"
        if business_type == 3:
            business_type_name = "Employee Benefits"
        if business_type == 4:
            business_type_name = "Funeral"
        if business_type == 5:
            business_type_name = "GAP Cover"
        if business_type == 6:
            business_type_name = "Recurring - Investment"
        if business_type == 7:
            business_type_name = "Lumpsum - Investment"
        if business_type == 8:
            business_type_name = "Investment- Both"
        if business_type == 9:
            business_type_name = "Medical Aid"
        if business_type == 10:
            business_type_name = "Other"
        if business_type == 11:
            business_type_name = "Will"
        if business_type == 12:
            business_type_name = "Risk"
        if business_type == 13:
            business_type_name = "ST Re-issued/instated"
        if business_type == 14:
            business_type_name = "Short Term Commercial"
        if business_type == 15:
            business_type_name = "Short Term Personal" 
        initiating_user_type = "Admin"
        if document.user.userType == 1:
            initiating_user_type = "ARC"
        if document.user.userType == 2:
            initiating_user_type = "Gatekeeper"
        if document.user.userType == 3:
            initiating_user_type = "Manager"
        if document.user.userType == 5:
            initiating_user_type = "B.A.C."
        gatekeeping_data = {}
        total_gatekeeping_versions = 0
        final_gatekeeping_score = 0
        gatekeeping = GateKeeping.objects.filter(document=document.pk)
        if gatekeeping.exists():
            total_gatekeeping_versions = gatekeeping.count()
            for counter, gatekeeping_record in enumerate(gatekeeping):
                fica = ""
                if gatekeeping_record.fica == 100:
                    fica = "Yes"
                elif gatekeeping_record.fica == 1:
                    fica = "No"
                elif gatekeeping_record.fica == 0:
                    fica = "N/A"
                proof_of_screening = ""
                if gatekeeping_record.proof_of_screening == 100:
                    proof_of_screening = "Yes"
                elif gatekeeping_record.proof_of_screening == 1:
                    proof_of_screening = "No"
                elif gatekeeping_record.proof_of_screening == 2:
                    proof_of_screening = "N/A"
                dra = ""
                if gatekeeping_record.dra == 100:
                    proof_of_screening = "Yes"
                elif gatekeeping_record.dra == 1:
                    proof_of_screening = "No"
                elif gatekeeping_record.dra == 0:
                    proof_of_screening = "N/A"
                letter_of_intro = ""
                if gatekeeping_record.letter_of_intro == 100:
                    letter_of_intro = "Yes"
                elif gatekeeping_record.letter_of_intro == 1:
                    letter_of_intro = "No"
                elif gatekeeping_record.letter_of_intro == 0:
                    letter_of_intro = "N/A"
                authorisation_letter = ""
                if gatekeeping_record.authorisation_letter == 100:
                    authorisation_letter = "Yes"
                elif gatekeeping_record.authorisation_letter == 1:
                    authorisation_letter = "No"
                elif gatekeeping_record.authorisation_letter == 0:
                    authorisation_letter = "N/A"
                roa_type = ""
                if gatekeeping_record.roa_type == 100:
                    roa_type = "SanFin ROA"
                elif gatekeeping_record.roa_type == 0:
                    roa_type = "SFP ROA"
                elif gatekeeping_record.roa_type == 1:
                    roa_type = "Glacier Ice"
                elif gatekeeping_record.roa_type == 2:
                    roa_type = "Compare Med"
                elif gatekeeping_record.roa_type == 3:
                    roa_type = "Get Quote"
                elif gatekeeping_record.roa_type == 4:
                    roa_type = "No"
                roa = ""
                if gatekeeping_record.roa == 100:
                    roa = "Yes"
                elif gatekeeping_record.roa == 1:
                    roa = "No"
                elif gatekeeping_record.roa == 0:
                    roa = "N/A"
                fna = ""
                if gatekeeping_record.fna == 100:
                    fna = "Yes"
                elif gatekeeping_record.fna == 1:
                    fna = "No"
                elif gatekeeping_record.fna == 0:
                    fna = "N/A"
                application = ""
                if gatekeeping_record.application == 100:
                    application = "Yes"
                elif gatekeeping_record.application == 1:
                    application = "No"
                elif gatekeeping_record.application == 0:
                    application = "N/A"
                quotation = ""
                if gatekeeping_record.quotation == 100:
                    quotation = "Yes"
                elif gatekeeping_record.quotation == 1:
                    quotation = "No"
                elif gatekeeping_record.quotation == 0:
                    quotation = "N/A"
                risk_portfolio = ""
                if gatekeeping_record.risk_portfolio == 1:
                    risk_portfolio = "Yes"
                elif gatekeeping_record.risk_portfolio == 1:
                    risk_portfolio = "No"
                elif gatekeeping_record.risk_portfolio == 0:
                    risk_portfolio = "N/A"
                mandate = ""
                if gatekeeping_record.mandate == 1:
                    mandate = "Yes"
                elif gatekeeping_record.mandate == 1:
                    mandate = "No"
                elif gatekeeping_record.mandate == 0:
                    mandate = "N/A"
                replacement = ""
                if gatekeeping_record.replacement == 1:
                    replacement = "Yes"
                elif gatekeeping_record.replacement == 1:
                    replacement = "No"
                elif gatekeeping_record.replacement == 0:
                    replacement = "N/A"
                replacement_type = ""
                if gatekeeping_record.replacement_type == 100:
                    replacement_type = "Rule 19"
                elif gatekeeping_record.replacement_type == 1:
                    replacement_type = "Not Rule 19"
                elif gatekeeping_record.replacement_type == 0:
                    replacement_type = "N/A"
                date_of_screening = ""
                if int(gatekeeping_record.date_of_screening) == 100:
                    date_of_screening = "Before quote date"
                elif int(gatekeeping_record.date_of_screening) == 1:
                    date_of_screening = "After quote date"
                elif gatekeeping_record.replacement_type == 0:
                    date_of_screening = "N/A"
                timeously = ""
                if int(gatekeeping_record.timeously) == 100:
                    timeously = "Yes"
                elif int(gatekeeping_record.timeously) == 1:
                    timeously = "No"
                elif int(gatekeeping_record.timeously) == 0:
                    timeously = "N/A"
                policy_schedule = ""
                if gatekeeping_record.policy_schedule == 100:
                    policy_schedule = "Yes"
                elif gatekeeping_record.policy_schedule == 1:
                    policy_schedule = "No"
                elif gatekeeping_record.policy_schedule == 0:
                    policy_schedule = "N/A"
                commission_release_form = ""
                if gatekeeping_record.commission_release_form == 100:
                    commission_release_form = "Yes"
                elif gatekeeping_record.commission_release_form == 1:
                    commission_release_form = "No"
                elif gatekeeping_record.commission_release_form == 0:
                    commission_release_form = "N/A"
                iteration = "First"
                if counter+1 == 1:
                    iteration = "First"
                elif counter+1 == 2:
                    iteration = "Second"
                elif counter+1 == 3:
                    iteration = "3rd"
                else:
                    iteration = f"{counter+1}th"
                total = 0
                score = 0
                if business_type == 1 or (business_type > 4 and business_type < 9) :
                    # score = gatekeepingDocument['fica'] + gatekeepingDocument['proof_of_screening'] + gatekeepingDocument['dra'] + gatekeepingDocument['letter_of_intro'] + gatekeepingDocument['authorisation_letter'] + gatekeepingDocument['roa_type'] + gatekeepingDocument['roa'] + gatekeepingDocument['fna'] + gatekeepingDocument['application'] + gatekeepingDocument['quotation'] + gatekeepingDocument['risk_portfolio'] + gatekeepingDocument['mandate'] + gatekeepingDocument['replacement'] + gatekeepingDocument['replacement_type']
                    gkDocument = GateKeeping.objects.filter(pk=gatekeeping_record.pk).values("fica","proof_of_screening","dra","letter_of_intro","authorisation_letter","roa","fna","application","quotation","risk_portfolio","mandate","replacement").latest('created_at')
                    for key in gkDocument:
                        if int(gkDocument[key]) == 100:
                            total += 100
                            score += int(gkDocument[key])
                        if int(gkDocument[key]) == 0:
                            total += 100
                    score = round(score/total *100) if total != 0 else 0
                if business_type == 3 or business_type == 4:
                    # score = gatekeepingDocument['fica'] + gatekeepingDocument['proof_of_screening'] + gatekeepingDocument['dra'] + gatekeepingDocument['letter_of_intro'] + gatekeepingDocument['authorisation_letter'] + gatekeepingDocument['roa_type'] + gatekeepingDocument['roa'] + gatekeepingDocument['fna'] + gatekeepingDocument['application'] + gatekeepingDocument['quotation'] + gatekeepingDocument['replacement'] + gatekeepingDocument['replacement_type']
                    gkDocument = GateKeeping.objects.filter(pk=gatekeeping_record.pk).values("fica","proof_of_screening","dra","letter_of_intro","authorisation_letter","roa","fna","application","quotation","replacement").latest('created_at')
                    for key in gkDocument:
                        if int(gkDocument[key]) == 100:
                            total += 100
                            score += int(gkDocument[key])
                        if int(gkDocument[key]) == 0:
                            total += 100
                    score = round(score/total *100) if total != 0 else 0
                if business_type == 5 or business_type == 9:
                    # score = gatekeepingDocument['fica'] + gatekeepingDocument['proof_of_screening'] + gatekeepingDocument['dra'] + gatekeepingDocument['letter_of_intro'] + gatekeepingDocument['authorisation_letter'] + gatekeepingDocument['roa_type'] + gatekeepingDocument['roa'] + gatekeepingDocument['fna'] + gatekeepingDocument['application'] + gatekeepingDocument['quotation']
                    gkDocument = GateKeeping.objects.filter(pk=gatekeeping_record.pk).values("fica","proof_of_screening","dra","letter_of_intro","authorisation_letter","roa","fna","application","quotation").latest('created_at')
                    for key in gkDocument:
                        if int(gkDocument[key]) == 100:
                            total += 100
                            score += int(gkDocument[key])
                        if int(gkDocument[key]) == 0:
                            total += 100
                    score = round(score/total *100) if total != 0 else 0
                if business_type == 12:
                    # score = gatekeepingDocument['fica'] + gatekeepingDocument['proof_of_screening'] + gatekeepingDocument['dra'] + gatekeepingDocument['letter_of_intro'] + gatekeepingDocument['authorisation_letter'] + gatekeepingDocument['roa_type'] + gatekeepingDocument['roa'] + gatekeepingDocument['fna'] + gatekeepingDocument['application'] + gatekeepingDocument['quotation'] + gatekeepingDocument['replacement'] + gatekeepingDocument['replacement_type']
                    gkDocument = GateKeeping.objects.filter(pk=gatekeeping_record.pk).values("fica","proof_of_screening","dra","letter_of_intro","authorisation_letter","roa","fna","application","quotation","replacement").latest('created_at')
                    for key in gkDocument:
                        if int(gkDocument[key]) == 100:
                            total += 100
                            score += int(gkDocument[key])
                        if int(gkDocument[key]) == 0:
                            total += 100
                    score = round(score/total *100) if total != 0 else 0
                if business_type == 10 :
                    # score = gatekeepingDocument['fica'] + gatekeepingDocument['proof_of_screening'] + gatekeepingDocument['dra'] + gatekeepingDocument['letter_of_intro'] + gatekeepingDocument['authorisation_letter'] + gatekeepingDocument['roa_type'] + gatekeepingDocument['roa']
                    gkDocument = GateKeeping.objects.filter(pk=gatekeeping_record.pk).values("fica","proof_of_screening","dra","letter_of_intro","authorisation_letter","roa").latest('created_at')
                    for key in gkDocument:
                        if int(gkDocument[key]) == 100:
                            total += 100
                            score += int(gkDocument[key])
                        if int(gkDocument[key]) == 0:
                            total += 100
                    score = round(score/total *100) if total != 0 else 0
                if business_type == 11 :
                    # score = gatekeepingDocument['fica'] + gatekeepingDocument['proof_of_screening'] + gatekeepingDocument['dra'] + gatekeepingDocument['letter_of_intro'] + gatekeepingDocument['authorisation_letter'] + gatekeepingDocument['application']
                    gkDocument = GateKeeping.objects.filter(pk=gatekeeping_record.pk).values("fica","proof_of_screening","dra","letter_of_intro","authorisation_letter","application").latest('created_at')
                    for key in gkDocument:
                        if int(gkDocument[key]) == 100:
                            total += 100
                            score += int(gkDocument[key])
                        if int(gkDocument[key]) == 0:
                            total += 100
                    score = round(score/total *100) if total != 0 else 0
                if business_type == 13 :
                    # score = gatekeepingDocument['fica'] + gatekeepingDocument['proof_of_screening'] + gatekeepingDocument['dra']
                    gkDocument = GateKeeping.objects.filter(pk=gatekeeping_record.pk).values("fica","policy_schedule","proof_of_screening","dra").latest('created_at')
                    for key in gkDocument:
                        if int(gkDocument[key]) == 100:
                            total += 100
                            score += int(gkDocument[key])
                        if int(gkDocument[key]) == 0:
                            total += 100
                    score = round(score/total *100) if total != 0 else 0
                if business_type == 14 or business_type == 15 :
                    # score = gatekeepingDocument['fica'] + gatekeepingDocument['proof_of_screening'] + gatekeepingDocument['dra'] + gatekeepingDocument['letter_of_intro'] + gatekeepingDocument['authorisation_letter'] + gatekeepingDocument['roa'] + gatekeepingDocument['fna'] + gatekeepingDocument['application'] + gatekeepingDocument['quotation'] + gatekeepingDocument['replacement']
                    gkDocument = GateKeeping.objects.filter(pk=gatekeeping_record.pk).values("fica","policy_schedule","proof_of_screening","dra","letter_of_intro","authorisation_letter","roa","fna","application","quotation","replacement").latest('created_at')
                    for key in gkDocument:
                        if int(gkDocument[key]) == 100:
                            total += 100
                            score += int(gkDocument[key])
                        if int(gkDocument[key]) == 0:
                            total += 100
                    score = round(score/total *100) if total != 0 else 0
                if business_type == 2 :
                # score = gatekeepingDocument['proof_of_screening']
                    gkDocument = GateKeeping.objects.filter(pk=gatekeeping_record.pk).values("commission_release_form" ).latest('created_at')
                    for key in gkDocument:
                        if int(gkDocument[key]) == 0:
                            total += 100
                        if int(gkDocument[key]) == 100:
                            total += 100
                            score += int(gkDocument[key])
                    score = round(score/total *100) if total != 0 else 0
                if counter+1 == total_gatekeeping_versions:
                    final_gatekeeping_score = score
                
                gatekeeping_data[f"{iteration} Review - Gatekeeper - FICA (Clear ID)"] = fica
                gatekeeping_data[f"{iteration} Review - Gatekeeper - Proof of Screening"] = proof_of_screening
                gatekeeping_data[f"{iteration} Review - Gatekeeper - DRA"] = dra
                gatekeeping_data[f"{iteration} Review - Gatekeeper - Letter of Introduction"] = letter_of_intro
                gatekeeping_data[f"{iteration} Review - Gatekeeper - Authorisation Letter"] = authorisation_letter
                gatekeeping_data[f"{iteration} Review - Gatekeeper - ROA Type"] = roa_type
                gatekeeping_data[f"{iteration} Review - Gatekeeper - ROA (All sections completed)"] = roa
                gatekeeping_data[f"{iteration} Review - Gatekeeper - FNA (Appropriate FNA filed"] = fna
                gatekeeping_data[f"{iteration} Review - Gatekeeper - Application"] = application
                gatekeeping_data[f"{iteration} Review - Gatekeeper - Quotation"] = quotation
                gatekeeping_data[f"{iteration} Review - Gatekeeper - Risk Portfolio"] = risk_portfolio
                gatekeeping_data[f"{iteration} Review - Gatekeeper - Mandate"] = mandate
                gatekeeping_data[f"{iteration} Review - Gatekeeper - Replacement"] = replacement
                gatekeeping_data[f"{iteration} Review - Gatekeeper - Replacement Type"] = replacement_type
                gatekeeping_data[f"{iteration} Review - Gatekeeper - Date of Screening"] = date_of_screening
                gatekeeping_data[f"{iteration} Review - Gatekeeper - Timeously"] = timeously
                gatekeeping_data[f"{iteration} Review - Gatekeeper - Policy Schedule"] = policy_schedule
                gatekeeping_data[f"{iteration} Review - Gatekeeper - Commission Release Form"] = commission_release_form
                gatekeeping_data[f"{iteration} Review - Gatekeeper - Score"] = score
        
        arc_data = {}
        total_arc_versions = 0
        final_arc_score = 0
        arc_records = arc.objects.filter(document=document.pk)
        if arc_records.exists():
            total_arc_versions = arc_records.count()
            for counter, arc_record in enumerate(arc_records):
                client_needs = ""
                if arc_record.client_needs == 100:
                    client_needs = "Approved"
                elif arc_record.client_needs == 1:
                    client_needs = "Partially Approved"
                elif arc_record.client_needs == 0:
                    client_needs = "Not Approved"
                appropriate_fna = ""
                if arc_record.appropriate_fna == 100:
                    appropriate_fna = "Approved"
                elif arc_record.appropriate_fna == 1:
                    appropriate_fna = "Partially Approved"
                elif arc_record.appropriate_fna == 0:
                    appropriate_fna = "Not Approved"
                fna_outcome = ""
                if arc_record.fna_outcome == 100:
                    fna_outcome = "Approved"
                elif arc_record.fna_outcome == 1:
                    fna_outcome = "Partially Approved"
                elif arc_record.fna_outcome == 0:
                    fna_outcome = "Not Approved"
                product_suitability = ""
                if arc_record.product_suitability == 100:
                    product_suitability = "Approved"
                elif arc_record.product_suitability == 1:
                    product_suitability = "Partially Approved"
                elif arc_record.product_suitability == 0:
                    product_suitability = "Not Approved"
                alternative_solutions = ""
                if arc_record.alternative_solutions == 100:
                    alternative_solutions = "Approved"
                elif arc_record.alternative_solutions == 1:
                    alternative_solutions = "Partially Approved"
                elif arc_record.alternative_solutions == 0:
                    alternative_solutions = "Not Approved"
                material_aspects = ""
                if arc_record.material_aspects == 100:
                    material_aspects = "Approved"
                elif arc_record.material_aspects == 1:
                    material_aspects = "Partially Approved"
                elif arc_record.material_aspects == 0:
                    material_aspects = "Not Approved"
                special_terms = ""
                if arc_record.special_terms == 100:
                    special_terms = "Approved"
                elif arc_record.special_terms == 1:
                    special_terms = "Partially Approved"
                elif arc_record.special_terms == 0:
                    special_terms = "Not Approved"
                replacement_terms = ""
                if arc_record.replacement_terms == 100:
                    replacement_terms = "Approved"
                elif arc_record.replacement_terms == 1:
                    replacement_terms = "Partially Approved"
                elif arc_record.replacement_terms == 0:
                    replacement_terms = "Not Approved"
                disclosure_a = ""
                if arc_record.disclosure_a == 100:
                    disclosure_a = "Approved"
                elif arc_record.disclosure_a == 1:
                    disclosure_a = "Partially Approved"
                elif arc_record.disclosure_a == 0:
                    disclosure_a = "Not Approved"
                disclosure_b = ""
                if arc_record.disclosure_b == 100:
                    disclosure_b = "Approved"
                elif arc_record.disclosure_b == 1:
                    disclosure_b = "Partially Approved"
                elif arc_record.disclosure_b == 0:
                    disclosure_b = "Not Approved"
                personal_details_a = ""
                if arc_record.personal_details_a == 100:
                    personal_details_a = "Approved"
                elif arc_record.personal_details_a == 1:
                    personal_details_a = "Partially Approved"
                elif arc_record.personal_details_a == 0:
                    personal_details_a = "Not Approved"
                general_a = ""
                if arc_record.general_a == 100:
                    general_a = "Approved"
                elif arc_record.general_a == 1:
                    general_a = "Partially Approved"
                elif arc_record.general_a == 0:
                    general_a = "Not Approved"
                general_b = ""
                if arc_record.general_b == 100:
                    general_b = "Approved"
                elif arc_record.general_b == 1:
                    general_b = "Partially Approved"
                elif arc_record.general_b == 0:
                    general_b = "Not Approved"
                general_c = ""
                if arc_record.general_c == 100:
                    general_c = "Approved"
                elif arc_record.general_c == 1:
                    general_c = "Partially Approved"
                elif arc_record.general_c == 0:
                    general_c = "Not Approved"
                general_d = ""
                if arc_record.general_d == 100:
                    general_d = "Approved"
                elif arc_record.general_d == 1:
                    general_d = "Partially Approved"
                elif arc_record.general_d == 0:
                    general_d = "Not Approved"
                risk_classes_a = ""
                if arc_record.risk_classes_a == 100:
                    risk_classes_a = "Approved"
                elif arc_record.risk_classes_a == 1:
                    risk_classes_a = "Partially Approved"
                elif arc_record.risk_classes_a == 0:
                    risk_classes_a = "Not Approved"
                risk_classes_b = ""
                if arc_record.risk_classes_b == 100:
                    risk_classes_b = "Approved"
                elif arc_record.risk_classes_b == 1:
                    risk_classes_b = "Partially Approved"
                elif arc_record.risk_classes_b == 0:
                    risk_classes_b = "Not Approved"
                fna_a = ""
                if arc_record.fna_a == 100:
                    fna_a = "Approved"
                elif arc_record.fna_a == 1:
                    fna_a = "Partially Approved"
                elif arc_record.fna_a == 0:
                    fna_a = "Not Approved"
                fna_b = ""
                if arc_record.fna_b == 100:
                    fna_b = "Approved"
                elif arc_record.fna_b == 1:
                    fna_b = "Partially Approved"
                elif arc_record.fna_b == 0:
                    fna_b = "Not Approved"
                recommended_products_a = ""
                if arc_record.recommended_products_a == 100:
                    recommended_products_a = "Approved"
                elif arc_record.recommended_products_a == 1:
                    recommended_products_a = "Partially Approved"
                elif arc_record.recommended_products_a == 0:
                    recommended_products_a = "Not Approved"
                recommended_products_b = ""
                if arc_record.recommended_products_b == 100:
                    recommended_products_b = "Approved"
                elif arc_record.recommended_products_b == 1:
                    recommended_products_b = "Partially Approved"
                elif arc_record.recommended_products_b == 0:
                    recommended_products_b = "Not Approved"
                recommended_products_c = ""
                if arc_record.recommended_products_c == 100:
                    recommended_products_c = "Approved"
                elif arc_record.recommended_products_c == 1:
                    recommended_products_c = "Partially Approved"
                elif arc_record.recommended_products_c == 0:
                    recommended_products_c = "Not Approved"
                replacements_a = ""
                if arc_record.replacements_a == 100:
                    replacements_a = "Approved"
                elif arc_record.replacements_a == 1:
                    replacements_a = "Partially Approved"
                elif arc_record.replacements_a == 0:
                    replacements_a = "Not Approved"
                replacements_b = ""
                if arc_record.replacements_b == 100:
                    replacements_b = "Approved"
                elif arc_record.replacements_b == 1:
                    replacements_b = "Partially Approved"
                elif arc_record.replacements_b == 0:
                    replacements_b = "Not Approved"
                replacements_c = ""
                if arc_record.replacements_c == 100:
                    replacements_c = "Approved"
                elif arc_record.replacements_c == 1:
                    replacements_c = "Partially Approved"
                elif arc_record.replacements_c == 0:
                    replacements_c = "Not Approved"
                replacements_d = ""
                if arc_record.replacements_d == 100:
                    replacements_d = "Approved"
                elif arc_record.replacements_d == 1:
                    replacements_d = "Partially Approved"
                elif arc_record.replacements_d == 0:
                    replacements_d = "Not Approved"
                client_consent_a = ""
                if arc_record.client_consent_a == 100:
                    client_consent_a = "Approved"
                elif arc_record.client_consent_a == 1:
                    client_consent_a = "Partially Approved"
                elif arc_record.client_consent_a == 0:
                    client_consent_a = "Not Approved"
                client_consent_b = ""
                if arc_record.client_consent_b == 100:
                    client_consent_b = "Approved"
                elif arc_record.client_consent_b == 1:
                    client_consent_b = "Partially Approved"
                elif arc_record.client_consent_b == 0:
                    client_consent_b = "Not Approved"
                arc_score = 0

                if business_type < 14 :
                    aDocument = arc.objects.filter(pk=arc_record.pk).values("client_needs","appropriate_fna","fna_outcome","product_suitability","alternative_solutions","material_aspects","special_terms","replacement_terms").latest('id')
                    arc_total = 120 if replacement == False else 100
                    for key in aDocument:
                        if key == "replacement_terms" and replacement:
                            continue    
                        arc_score += aDocument[key]
                    arc_score = round(arc_score/arc_total *100) if total != 0 else 0
                if business_type >= 14 :
                    aDocument = arc.objects.filter(pk=arc_record.pk).values("disclosure_a", "disclosure_b", "personal_details_a", "personal_details_b", "general_a", "general_b", "general_c", "general_d", "risk_classes_a", "risk_classes_b", "fna_a", "fna_b", "recommended_products_a", "recommended_products_b", "recommended_products_c", "replacements_a", "replacements_b", "replacements_c", "replacements_d", "client_consent_a", "client_consent_b", ).latest('id')
                    for key in aDocument:
                        arc_score += aDocument[key]
                    arc_total = 100
                    for key in aDocument:
                        if aDocument[key] == 5:
                            arc_score += aDocument[key]
                    arc_score = round(arc_score/arc_total *100) if total != 0 else 0

                if counter+1 == total_arc_versions:
                    final_arc_score = arc_score
                
                iteration = "First"
                if counter+1 == 1:
                    iteration = "First"
                elif counter+1 == 2:
                    iteration = "Second"
                elif counter+1 == 3:
                    iteration = "3rd"
                else:
                    iteration = f"{counter+1}th"
                
                arc_data[f"{iteration} Review - ARC - Clients needs"] = client_needs
                arc_data[f"{iteration} Review - ARC - Appropriate FNA"] = appropriate_fna
                arc_data[f"{iteration} Review - ARC - FNA outcome"] = fna_outcome
                arc_data[f"{iteration} Review - ARC - Product suitability"] = product_suitability
                arc_data[f"{iteration} Review - ARC - Alternative Solutions"] = alternative_solutions
                arc_data[f"{iteration} Review - ARC - Material Aspects"] = material_aspects
                arc_data[f"{iteration} Review - ARC - Special Terms"] = special_terms
                arc_data[f"{iteration} Review - ARC - Replacement Terms"] = replacement_terms
                arc_data[f"{iteration} Review - ARC - Disclosures Questions: Has FAIS disclosure documents (permit) been signed and filed?"] = disclosure_a
                arc_data[f"{iteration} Review - ARC - Disclosures Questions: Has client has provided authority to access information?"] = disclosure_b
                arc_data[f"{iteration} Review - ARC - General Questions: Personal Details Questions: Was the name, surname gender, identifying number and client preference filled in correctly?"] = personal_details_a
                arc_data[f"{iteration} Review - ARC - General Questions: Has the refusal of cover question been answered and details provided where the response is positive?"] = general_a
                arc_data[f"{iteration} Review - ARC - General Questions: Is this a replacement (or partial replacement) of cover?"] = general_b
                arc_data[f"{iteration} Review - ARC - General Questions: List of clients’ previous insurer"] = general_c
                arc_data[f"{iteration} Review - ARC - General Questions: Has details of previous losses been completed with details?"] = general_d
                arc_data[f"{iteration} Review - ARC - Risk Classes / Client needs (Factual Questions): Has all the relevant risk classes been identified and offered?"] = risk_classes_a
                arc_data[f"{iteration} Review - ARC - Risk Classes / Client needs (Factual Questions): Have all questions in the identified risk classes been answered"] = risk_classes_b
                arc_data[f"{iteration} Review - ARC - Financial Needs Analysis Summary / Cover Recommendations: Was the need of the client identified in the ROA?"] = fna_a
                arc_data[f"{iteration} Review - ARC - Financial Needs Analysis Summary / Cover Recommendations: Was option or additional cover offered and discussed with the client? Are the responses recorded?"] = fna_b
                arc_data[f"{iteration} Review - ARC - Recommended Products: Has the optional cover recommendations been made for all risk classes? Optional cover on buildings, contents, all risk, rental car, caravan contents or glitter finish on watercraft, etc. "] = recommended_products_a
                arc_data[f"{iteration} Review - ARC - Recommended Products: Is there an explanation that illustrates exclusions, special conditions or endorsements, or a reason to refuse cover? (e.g. no burglar bars, or linked alarm, immobilizer)"] = recommended_products_b
                arc_data[f"{iteration} Review - ARC - Recommended Products: Is the quote on file?"] = recommended_products_c
                arc_data[f"{iteration} Review - ARC - Replacements: Is this product a replacement? "] = replacements_a
                arc_data[f"{iteration} Review - ARC - Replacements: Has the purpose and reasons been identified?"] = replacements_b
                arc_data[f"{iteration} Review - ARC - Replacements: Has the comparison section been completed?"] = replacements_c
                arc_data[f"{iteration} Review - ARC - Replacements: Has the required information been captured adequately on the Comparison ROA?"] = replacements_d
                arc_data[f"{iteration} Review - ARC - Client Consent: Has the debit order instruction been completed and signed?"] = client_consent_a
                arc_data[f"{iteration} Review - ARC - Client Consent: Has the client consented to a credit check for STI purposes?"] = client_consent_b
                arc_data[f"{iteration} Review - ARC - Score"] = arc_score
        comments_data = {}
        comments = DocumentComments.objects.filter(document=document.pk)
        for counter, comment in enumerate(comments):
            comments_data[f"Comment {counter+1} By"] = comment.user.first_name
            comments_data[f"Comment {counter+1}"] = comment.comment
        picked_up_name = ""
        picked_up_email = ""
        picked_up_type = ""
        if document.picked_up != 0:
            picked_up_data = UserAccount.objects.filter(pk=document.picked_up)
            if picked_up_data.exists():
                picked_up_data = picked_up_data.first()
                picked_up_name = picked_up_data.first_name + " " + picked_up_data.last_name
                picked_up_email = picked_up_data.email
                picked_up_type = "ARC" if picked_up_data.userType == 1 else "Gatekeeper"
                picked_up_user_profile = user_profile.objects.filter(user=picked_up_data.pk)
                if picked_up_user_profile.exists():
                    picked_up_user_profile = picked_up_user_profile.first()
                    picked_up_name = picked_up_user_profile.Full_Name
        document_data = {
            "Document Creation Date" : document.created_at.strftime('%H:%M:%S %d %b %Y'), 
            "Document Last Update Date" : document.updated_at.strftime('%H:%M:%S %d %b %Y'), 
            "Policy Number" : document.policy_number, 
            "Initiating User" : document.user.first_name + " " + document.user.last_name, 
            "Initiating User Email" : document.user.email, 
            "Initiating User Type" : "ARC" if initiating_user_type == 1 else "Gatekeeper", 
            "Picked Up by User" : picked_up_name, 
            "Picked Up by User Email" : picked_up_email, 
            "Picked Up by User Type" : picked_up_type, 
            "Advisor" : advisor_name, 
            "Advisor Email" : document.advisor.email, 
            "Advisor ID Number" : advisor_id, 
            "Advisor Region." : region_name, 
            "Advisor Regional Manager" : regional_manager_name, 
            "Advisor Regional Manager Email" : regional_manager_email, 
            "Advisor B.A.C." : bac_name, 
            "Advisor B.A.C. Email" : bac_email, 
            "Client Name" : document.clientName, 
            "Supplier" : document.supplier, 
            "Product" : document.product, 
            "Business Type" : business_type_name, 
            "Starting Type" : "Gatekeeper Level" if document.starting_point == 1 else "ARC Level", 
            "Gatekeeping Versions" : total_gatekeeping_versions, 
            "Gatekeeping Final Score" : final_gatekeeping_score, 
            "ARC Versions" : total_arc_versions, 
            "ARC Final Score" : final_arc_score, 
            "Lump Sum" : str(document.lump_sum).replace(".",","), 
            "Monthly Premium" : str(document.monthly_premium).replace(".",","), 
            "Commission" : str(document.commission).replace(".",","), 
            # "Lump Sum" : babel.numbers.format_currency(str(document.lump_sum).replace(" ","").replace(",","."), 'ZAR', locale='en_ZA', currency_digits=False), 
            # "Monthly Premium" : babel.numbers.format_currency(str(document.monthly_premium).replace(" ","").replace(",","."), 'ZAR', locale='en_ZA', currency_digits=False), 
            # "Commission" : babel.numbers.format_currency(str(document.commission).replace(" ","").replace(",","."), 'ZAR', locale='en_ZA', currency_digits=False), 
        }
        document_data.update(gatekeeping_data)
        document_data.update(arc_data)
        document_data.update(comments_data)
        requestedData.append(document_data)
    document_df = pd.DataFrame(requestedData)
    if len(requestedData) > 0:
        file_name = f"static/csv/{file_initial_name}_complaince_export_{uuid.uuid4()}.csv"
        document_df.to_csv(f"{file_name}")
        notificationData = {
            "account" : user.pk,
            "notificationType" : 6,        
            "user" : user.pk,
            "title" : title,
            "message" : f"{message} is ready.",
            "downloading_link" : file_name,
            "status" : False
        }
        serializer = NotificationsSerializer(data=notificationData)
        if serializer.is_valid():
            serializer.create(serializer.validated_data)
        else:
            logger.info(serializer.errors)
        user_name = user.first_name + " " + user.last_name
        userData = user_profile.objects.filter(user=user.pk)
        if userData.exists():
            user_name = userData.first().Full_Name
        downloading_link = env('DJANGO_BACKEND_URL') + "/" + file_name
        notificationData = {
            "account" : user.pk,
            "notificationType" : 2,        
            "user" : user.pk,
            "title" : title,
            "message" : f"""
                <div style="max-width: 600px; margin: 20px auto; padding: 20px; background-color: #fff; border-radius: 8px; box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);">
                    <h3 class='text-center' style="color: #333; margin-bottom: 20px;">{title}</h3>
                    <div style="background-color: #FFCC00; color: #333; padding: 15px; border-radius: 4px; margin-bottom: 20px;">
                        <p style="margin: 0;">Dear {user_name},</p>
                        <p style="margin: 0;">Your requested compliance export is ready to be downloaded.</p>
                        <div class='text-center'>
                            <a id="downloading_link" href="{downloading_link}" target="_blank" class="btn btn-primary btn-sm btn-sfp">Download</a>
                        </div>
                    </div>
                    <p style="color: #666; line-height: 1.6; margin-bottom: 15px;">If you have any questions or concerns feel free to reach out.</p>
                    <p style="color: #666; line-height: 1.6; margin-bottom: 15px;">Best regards,<br>Succession</p>
                </div>
            """,
            "status" : False,
        }
        serializer = NotificationsSerializer(data=notificationData)
        if serializer.is_valid():
            serializer.create(serializer.validated_data)
        else:
            logger.info(serializer.errors)
        
    else:
        notificationData = {
            "account" : user.pk,
            "notificationType" : 6,        
            "user" : user.pk,
            "title" : title,
            "message" : "No data found in given range",
            "status" : False,
        }
        serializer = NotificationsSerializer(data=notificationData)
        if serializer.is_valid():
            serializer.create(serializer.validated_data)
        else:
            logger.info(serializer.errors)
        

@shared_task
def roa_disclosure_products_update(userId, data):
    input_user = UserAccount.objects.get(id=userId)
    title = "Disclosure Products"
    message  = "ROA Disclosures Products bulk update started"
    notificationData = {
        "account" : input_user.pk,
        "notificationType" : 6,        
        "user" : input_user.pk,
        "title" : title,
        "message" : message,
        "status" : False,
    }
    serializer = NotificationsSerializer(data=notificationData)
    if serializer.is_valid():
        serializer.create(serializer.validated_data)
    else:
        logger.info(serializer.errors)
    
    csv_data = data['productsCsvFile']
    format, csvstr = csv_data.split(';base64,')
    ext = format.split('/')[-1]
    file_name = "'productsCsvFile." + ext
    csvData = ContentFile(base64.b64decode(csvstr), name=file_name) 
    sheets = pd.read_excel(csvData, sheet_name=None)

    # Access sheet names as a dictionary
    sheet_names = sheets.keys()

    logData = {
        "account" : input_user.pk,
        "log_name" : f"Bulk Products Upload {uuid.uuid4()}",
        "log_type" : 1,
        "status" : 0,
    }
    log_serializer = LogSerializer(data=logData)
    if log_serializer.is_valid():
        log_serializer.save()
        log_id = log_serializer.data['id']
    else:
        logger.info(log_serializer.errors)
    logKPIs = {
        "account" : input_user.pk,
        "log" : log_id,
        "kpis" : {
            "total_users": 0,"total_products":0, "total_products_added": 0, "total_user_products_updated" : 0,
        }
    }
    log_kpis_serializer = LogKPIsSerializer(data=logKPIs)
    if log_kpis_serializer.is_valid():
        log_kpis_serializer.create(log_kpis_serializer.validated_data)
    else:
        logger.info(log_kpis_serializer.errors)
    
    
    total_users = 0
    total_sheets = 0
    total_products = 0
    total_products_added = 0
    total_products_updated = 0
    for sheet in sheet_names:
        logger.info(f"Started for {sheet}")
        total_sheets += 1
        logContent = {
            "account" : input_user.pk,
            "log" : log_id,
            "log_type" : 1,
            "log_description" : f"Updating products for {sheet}.",
        } 
        log_content_serializer = LogContentSerializer(data=logContent)
        if log_content_serializer.is_valid():
            log_content_serializer.create(log_content_serializer.validated_data)
        else:
            logger.info(log_content_serializer.errors)
        
        product_type = 1
        if "ST" in sheet:
            product_type = 1
        if "Health" in sheet:
            product_type = 2
        if "Life" in sheet:
            product_type = 3

        disclosures_product_df = pd.read_excel(csvData, sheet_name=sheet, header = 1)

        disclosures_product_df.fillna('', inplace=True)
        disclosures_product_df.rename(columns={'Unnamed: 0': 'ProductsList'}, inplace=True)
        # disclosures_product_df.columns = disclosures_product_df.columns.str.replace(' ', '_').str.replace('__', '_').str.replace('.', '_')

        logContent = {
            "account" : input_user.pk,
            "log" : log_id,
            "log_type" : 1,
            "log_description" : f"Updating products in database.",
        } 
        log_content_serializer = LogContentSerializer(data=logContent)
        if log_content_serializer.is_valid():
            log_content_serializer.create(log_content_serializer.validated_data)
        else:
            logger.info(log_content_serializer.errors)
        total_users = 0
        users = disclosures_product_df['ProductsList'].tolist()
        for user in users:
            if user == "" or user == "Email" or user == "N/A":
                continue
            user_product_data = disclosures_product_df.iloc[disclosures_product_df.loc[disclosures_product_df['ProductsList']==user].index[0]].to_dict()
            update_user_product_data = {k: v for k, v in user_product_data.items() if v != ""}
            user_data = user_profile.objects.filter(user__email__iexact=user)
            if user_data.exists():
                user_data = user_data.first()
                for product in user_product_data.keys():
                    if product == "MPY":
                        continue
                    if product == "Sanlam":
                        continue
                    if product == "ProductsList":
                        continue
                    if "Unnamed" in product:
                        continue
                    disclosure_product = DisclosuresProductProviders.objects.filter(product=product)
                    if disclosure_product.exists():
                        total_products_updated += 1
                        update_log = f"<p>Product {product} exists and updated.</p>"
                        disclosure_product_data = {
                            "product" : product,
                            "product_type" : product_type,
                        }
                        disclosure_product_serializer = DisclosuresProductProviders_Serializer(instance=disclosure_product.first(), data=disclosure_product_data, partial=True)
                        if disclosure_product_serializer.is_valid():
                            disclosure_product_serializer.save()
                            logContent = {
                                "account" : input_user.pk,
                                "log" : log_id,
                                "log_type" : 1,
                                "log_description" : f"Product {product} exists and updated.",
                            } 
                            log_content_serializer = LogContentSerializer(data=logContent)
                            if log_content_serializer.is_valid():
                                log_content_serializer.create(log_content_serializer.validated_data)
                            else:
                                print(log_content_serializer.errors)
                            update_user_product_data = {
                                "email": user,
                                "user": user_data.user.pk,
                                "product": disclosure_product.first().pk,
                                "subcode": user_product_data[product]
                            }
                            old_data = DisclosuresAdvisorSubCodes.objects.filter(user=user_data.user.pk, product=disclosure_product.first().pk)
                            if old_data.exists():
                                serializer = DisclosuresAdvisorSubCodes_Serializer(instance=old_data.first(), data=update_user_product_data)
                                if serializer.is_valid():
                                    serializer.save()
                                    logContent = {
                                        "account" : input_user.pk,
                                        "log" : log_id,
                                        "log_type" : 1,
                                        "log_description" : f"Product {product} for user {user_data.Full_Name} updated.",
                                    } 
                                    log_content_serializer = LogContentSerializer(data=logContent)
                                    if log_content_serializer.is_valid():
                                        log_content_serializer.create(log_content_serializer.validated_data)
                                    else:
                                        print(log_content_serializer.errors)
                            else:
                                serializer = DisclosuresAdvisorSubCodes_Serializer(data = update_user_product_data)
                                if serializer.is_valid():
                                    serializer.create(serializer.validated_data)
                                    logContent = {
                                        "account" : input_user.pk,
                                        "log" : log_id,
                                        "log_type" : 1,
                                        "log_description" : f"Product {product} for user {user_data.Full_Name} added.",
                                    } 
                                    log_content_serializer = LogContentSerializer(data=logContent)
                                    if log_content_serializer.is_valid():
                                        log_content_serializer.create(log_content_serializer.validated_data)
                                    else:
                                        print(log_content_serializer.errors)
                        else:
                            print(disclosure_product_serializer.errors)
                            total_products_added += 1
                    else:
                        disclosure_product_data = {
                            "product" : product,
                            "product_type" : product_type,
                        }
                        disclosure_product_serializer = DisclosuresProductProviders_Serializer(data=disclosure_product_data)
                        if disclosure_product_serializer.is_valid():
                            disclosure_product = disclosure_product_serializer.create(disclosure_product_serializer.validated_data)
                            logContent = {
                                "account" : input_user.pk,
                                "log" : log_id,
                                "log_type" : 1,
                                "log_description" : f"Product {product} already exists.",
                            } 
                            log_content_serializer = LogContentSerializer(data=logContent)
                            if log_content_serializer.is_valid():
                                log_content_serializer.create(log_content_serializer.validated_data)
                            else:
                                print(log_content_serializer.errors)
                        else:
                            print(disclosure_product_serializer.errors) 
                        update_user_product_data = {
                            "email": user,
                            "user": user_data.user.pk,
                            "product": disclosure_product.pk,
                            "subcode": user_product_data[product]
                        }
                        old_data = DisclosuresAdvisorSubCodes.objects.filter(user=user_data.user.pk, product=disclosure_product.pk)
                        if old_data.exists():
                            serializer = DisclosuresAdvisorSubCodes_Serializer(instance=old_data.first(), data = update_user_product_data)
                            if serializer.is_valid():
                                serializer.save()
                                logContent = {
                                    "account" : input_user.pk,
                                    "log" : log_id,
                                    "log_type" : 1,
                                    "log_description" : f"Product {product} for user {user_data.Full_Name} updated.",
                                } 
                                log_content_serializer = LogContentSerializer(data=logContent)
                                if log_content_serializer.is_valid():
                                    log_content_serializer.create(log_content_serializer.validated_data)
                                else:
                                    print(log_content_serializer.errors)
                        else:
                            serializer = DisclosuresAdvisorSubCodes_Serializer(data = update_user_product_data)
                            if serializer.is_valid():
                                serializer.create(serializer.validated_data)
                                logContent = {
                                    "account" : input_user.pk,
                                    "log" : log_id,
                                    "log_type" : 1,
                                    "log_description" : f"Product {product} for user {user_data.Full_Name} added.",
                                } 
                                log_content_serializer = LogContentSerializer(data=logContent)
                                if log_content_serializer.is_valid():
                                    log_content_serializer.create(log_content_serializer.validated_data)
                                else:
                                    print(log_content_serializer.errors)
    title = "Disclosure Products"
    message  = "ROA Disclosures Products bulk update completed"
    notificationData = {
        "account" : input_user.pk,
        "notificationType" : 6,        
        "user" : input_user.pk,
        "title" : title,
        "message" : message,
        "log" : log_id,
        "status" : False,
    }
    serializer = NotificationsSerializer(data=notificationData)
    if serializer.is_valid():
        serializer.create(serializer.validated_data)
    else:
        logger.info(serializer.errors)    
    user_name = input_user.first_name + " " + input_user.last_name
    user_profile_data = user_profile.objects.filter(user=input_user.pk)
    if user_profile_data.exists():
        user_profile_data = user_profile_data.first()
        user_name = user_profile_data.Full_Name
    notificationData = {
        "account" : input_user.pk,
        "notificationType" : 2,        
        "user" : input_user.pk,
        "title" : title,
        "message" : f"""
            <div style="max-width: 600px; margin: 20px auto; padding: 20px; background-color: #fff; border-radius: 8px; box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);">
                <h3 class='text-center' style="color: #333; margin-bottom: 20px;">{title}</h3>
                <div style="background-color: #FFCC00; color: #333; padding: 15px; border-radius: 4px; margin-bottom: 20px;">
                    <p style="margin: 0;">Dear {user_name},</p>
                    <p style="margin: 0;">Your requested disclosures product update now completed.</p>
                </div>
                <p style="color: #666; line-height: 1.6; margin-bottom: 15px;">If you have any questions or concerns feel free to reach out.</p>
                <p style="color: #666; line-height: 1.6; margin-bottom: 15px;">Best regards,<br>Succession</p>
            </div>
        """,
        "status" : False,
    }
    serializer = NotificationsSerializer(data=notificationData)
    if serializer.is_valid():
        serializer.create(serializer.validated_data)
    else:
        logger.info(serializer.errors)
    logger.info(f"Completed at {datetime.now()}")

@shared_task
def commission_report(userId, requestedData):
    user = UserAccount.objects.get(id=userId)
    title = "Commission Report Export"
    message = "Commission Report complete data export"
    notificationData = {
        "account" : user.pk,
        "notificationType" : 6,        
        "user" : user.pk,
        "title" : title,
        "message" : message,
        "status" : False,
    }
    serializer = NotificationsSerializer(data=notificationData)
    if serializer.is_valid():
        serializer.create(serializer.validated_data)
    else:
        logger.info(serializer.errors)
    filterType = int(requestedData['filterType'])
    year = requestedData['year']
    monthyear = requestedData['monthyear']
    month = requestedData['month']
    date = requestedData['date']
    fromdate = requestedData['fromdate']
    todate = requestedData['todate']
    customFilterType = int(requestedData['customFilterType'])
    region = (requestedData['region'])
    advisor = (requestedData['advisor'])
    businessType = (requestedData['businessType'])
    # Annual Data
    reviewsData = ComplianceDocument.objects.all()
    file_initial_name = "Report for all"
    if not user.is_superuser:
        if user.userType == 1:
            reviewsData = reviewsData.filter(user=user.pk)
        if user.userType == 2:
            reviewsData = reviewsData.filter(user=user.pk)
        if user.userType == 3:
            regional_manager = region_manager.objects.filter(manager=user.pk)
            if regional_manager.exists():
                advisor_ids = user_profile.objects.filter(region=regional_manager.first().region.pk)
                if advisor_ids.exists():
                    advisor_ids = list(advisor_ids.values_list('user',flat=True))
                    reviewsData = reviewsData.filter(advisor__in=advisor_ids)
                reviewsData = reviewsData.filter(advisor__in=advisor_ids)
        if user.userType == 5:
            advisor_ids = user_profile.objects.filter(bac=user.pk)
            if advisor_ids.exists():
                advisor_ids = list(advisor_ids.values_list('user',flat=True))
                reviewsData = reviewsData.filter(advisor__in=advisor_ids)
        if user.userType == 6:
            reviewsData = reviewsData.filter(advisor=user.pk)
    if filterType == 1:
        reviewsData = reviewsData.filter(updated_at__year=year)
        file_initial_name = f"Report for year {year}"
    if filterType == 2:
        reviewsData = reviewsData.filter(updated_at__year=monthyear, updated_at__month=month)
        date_range_ = datetime.strptime(f"{monthyear}-{month}", "%Y-%m").strftime('%b %Y')
        file_initial_name = f"Report for {date_range_}"
    if filterType == 3:
        reviewsData = reviewsData.filter(updated_at__date=date)
        date_range_ = datetime.strptime(f"{date}", "%Y-%m-%d").strftime('%d %b %Y')
        file_initial_name = f"Report for {date_range_}"
    if filterType == 4:
        date_range = (datetime.strptime(fromdate, '%Y-%m-%d') , datetime.strptime(todate, '%Y-%m-%d') + timedelta(days=1))
        reviewsData = reviewsData.filter(updated_at__range=date_range)
        date_range_ = "From" + datetime.strptime(f"{fromdate}", "%Y-%m-%d").strftime('%d %b %Y') + " to " + datetime.strptime(f"{todate}", "%Y-%m-%d").strftime('%d %b %Y')
        file_initial_name = f"Report for date range {date_range_}"
    if region != "all":
        reviewsData = reviewsData.filter(region=region)
    if advisor != "all":
        reviewsData = reviewsData.filter(advisor=int(advisor))
    if businessType != "all":
        reviewsData = reviewsData.filter(businessType=int(businessType))
    
    total_reviews = reviewsData.count()
    total_documents = reviewsData.values()
    total_regions = reviewsData.values('region').distinct().count()
    total_advisors = reviewsData.values('advisor').distinct().count()
    total_commission = reviewsData.aggregate(total_commission=Sum(Cast('commission', output_field=FloatField())))['total_commission']
    total_commission = round(total_commission,2) if total_commission else 0.0
    # for review_document in total_documents:
    #     gk = GateKeeping.objects.filter(document=review_document['id'])
    #     if gk.exists():
    #         gk = gk.values().latest('version')
    #         total_commission += float(gk['commission'].replace(',', '.'))
    # Date wise Trend
    commission_trend = {}
    labels = []
    values = []
    if filterType == 1:
        datewise_data = reviewsData.values('updated_at__year','updated_at__month').distinct().order_by('updated_at__year','updated_at__month')
        for date in datewise_data:
            commission = reviewsData.filter(updated_at__year=date['updated_at__year'], updated_at__month=date['updated_at__month'])
            if commission.exists():
                commission = commission.aggregate(total_commission=Sum(Cast('commission', output_field=FloatField())))['total_commission']
            else:
                commission = 0
                    # commission_trend.append({"date" : review_document['updated_at__date'].strftime('%d %b %Y'), "commission": float(gk['commission'].replace(',', '.'))})
            labels.append(datetime.strftime(datetime.strptime(f"{date['updated_at__year']}-{date['updated_at__month']}", '%Y-%m') , '%b %Y'))
            values.append(int(commission))
    if filterType == 2:
        datewise_data = reviewsData.values('updated_at__year','updated_at__month', 'updated_at__day').distinct().order_by('updated_at__year','updated_at__month', 'updated_at__day')
        for date in datewise_data:
            commission = reviewsData.filter(updated_at__year=date['updated_at__year'], updated_at__month=date['updated_at__month'], updated_at__day=date['updated_at__day'])
            if commission.exists():
                commission = commission.aggregate(total_commission=Sum(Cast('commission', output_field=FloatField())))['total_commission']
            else:
                commission = 0
                    # commission_trend.append({"date" : review_document['updated_at__date'].strftime('%d %b %Y'), "commission": float(gk['commission'].replace(',', '.'))})
            labels.append(datetime.strftime(datetime.strptime(f"{date['updated_at__year']}-{date['updated_at__month']}-{date['updated_at__day']}", '%Y-%m-%d') , '%d %b %Y'))
            values.append(int(commission))
    if filterType == 3:
        datewise_data = reviewsData.values('updated_at__date', 'updated_at__hour').distinct().order_by('updated_at__date', 'updated_at__hour')
        for date in datewise_data:
            commission = reviewsData.filter(updated_at__date=date['updated_at__date'], updated_at__hour=date['updated_at__hour'])
            if commission.exists():
                commission = commission.aggregate(total_commission=Sum(Cast('commission', output_field=FloatField())))['total_commission']
            else:
                commission = 0
                    # commission_trend.append({"date" : review_document['updated_at__date'].strftime('%d %b %Y'), "commission": float(gk['commission'].replace(',', '.'))})
            labels.append(datetime.strftime(datetime.strptime(f"{date['updated_at__date']} {date['updated_at__hour']}", '%Y-%m-%d %H'), "%I %p"))
            values.append(int(commission))
    if filterType == 4:
        if customFilterType == 1:
            if (datetime.strptime(todate, "%Y-%m-%d") - datetime.strptime(fromdate, "%Y-%m-%d")).days > 30:
                datewise_data = reviewsData.values('updated_at__year','updated_at__month').distinct().order_by('updated_at__year','updated_at__month')
                for date in datewise_data:
                    commission = reviewsData.filter(updated_at__year=date['updated_at__year'], updated_at__month=date['updated_at__month'])
                    if commission.exists():
                        commission = commission.aggregate(total_commission=Sum(Cast('commission', output_field=FloatField())))['total_commission']
                    else:
                        commission = 0
                            # commission_trend.append({"date" : review_document['updated_at__date'].strftime('%d %b %Y'), "commission": float(gk['commission'].replace(',', '.'))})
                    labels.append(datetime.strftime(datetime.strptime(f"{date['updated_at__year']}-{date['updated_at__month']}", '%Y-%m') , '%b %Y'))
                    values.append(int(commission))
            else:
                datewise_data = reviewsData.values('updated_at__date').distinct().order_by('updated_at__date')
                for date in datewise_data:
                    commission = reviewsData.filter(updated_at__date=date['updated_at__date'])
                    if commission.exists():
                        commission = commission.values('updated_at__date').aggregate(total_commission=Sum(Cast('commission', output_field=FloatField())))['total_commission']
                    else:
                        commission = 0
                            # commission_trend.append({"date" : review_document['updated_at__date'].strftime('%d %b %Y'), "commission": float(gk['commission'].replace(',', '.'))})
                    labels.append(date['updated_at__date'].strftime('%d %b %Y'))
                    values.append(int(commission))
        if customFilterType == 2:
            datewise_data = reviewsData.values('updated_at__year','updated_at__week').distinct().order_by('updated_at__year','updated_at__week')
            for date in datewise_data:
                commission = reviewsData.filter(updated_at__year=date['updated_at__year'], updated_at__week=date['updated_at__week'])
                if commission.exists():
                    commission = commission.aggregate(total_commission=Sum(Cast('commission', output_field=FloatField())))['total_commission']
                else:
                    commission = 0
                        # commission_trend.append({"date" : review_document['updated_at__date'].strftime('%d %b %Y'), "commission": float(gk['commission'].replace(',', '.'))})
                labels.append(f"{date['updated_at__year']} Week {date['updated_at__week']}")
                values.append(int(commission))
        if customFilterType == 3:
            datewise_data = reviewsData.values('updated_at__year','updated_at__month').distinct().order_by('updated_at__year','updated_at__month')
            for date in datewise_data:
                commission = reviewsData.filter(updated_at__year=date['updated_at__year'], updated_at__month=date['updated_at__month'])
                if commission.exists():
                    commission = commission.aggregate(total_commission=Sum(Cast('commission', output_field=FloatField())))['total_commission']
                else:
                    commission = 0
                        # commission_trend.append({"date" : review_document['updated_at__date'].strftime('%d %b %Y'), "commission": float(gk['commission'].replace(',', '.'))})
                labels.append(datetime.strftime(datetime.strptime(f"{date['updated_at__year']}-{date['updated_at__month']}", '%Y-%m') , '%b %Y'))
                values.append(int(commission))
        if customFilterType == 4:
            datewise_data = reviewsData.values('updated_at__year','updated_at__quarter').distinct().order_by('updated_at__year','updated_at__quarter')
            for date in datewise_data:
                commission = reviewsData.filter(updated_at__year=date['updated_at__year'], updated_at__quarter=date['updated_at__quarter'])
                if commission.exists():
                    commission = commission.aggregate(total_commission=Sum(Cast('commission', output_field=FloatField())))['total_commission']
                else:
                    commission = 0
                        # commission_trend.append({"date" : review_document['updated_at__date'].strftime('%d %b %Y'), "commission": float(gk['commission'].replace(',', '.'))})
                labels.append(f"{date['updated_at__year']} Quarter {date['updated_at__quarter']}")
                values.append(int(commission))
        if customFilterType == 5:
            datewise_data = reviewsData.values('updated_at__year').distinct().order_by('updated_at__year')
            for date in datewise_data:
                commission = reviewsData.filter(updated_at__year=date['updated_at__year'])
                if commission.exists():
                    commission = commission.aggregate(total_commission=Sum(Cast('commission', output_field=FloatField())))['total_commission']
                else:
                    commission = 0
                labels.append(f"{date['updated_at__year']}")
                values.append(int(commission))
    commission_trend = {
        "labels" : labels,
        "values" : values
    }
    # Regions
    available_regions = regions.objects.all().values('region')
    if user.userType == 3:
        regional_manager = region_manager.objects.filter(manager=user.pk)
        if regional_manager.exists():
            available_regions = available_regions.filter(id=regional_manager.first().region.pk).values('region')
    if user.userType == 5:
        region_ids = user_profile.objects.filter(bac=user.pk)
        if region_ids.exists():
            region_ids = list(region_ids.values_list('region',flat=True))
            available_regions = available_regions.filter(id__in=region_ids).values('region')
    if user.userType == 6:
        region = user_profile.objects.filter(user=user.pk).first().region
        reviewsData = reviewsData.filter(region=region.region)
        available_regions = available_regions.filter(id=region.pk).values('region')
    # Region wise Trend
    region_commission_trend = {}
    regionsData = []
    top_regions = []
    labels = []
    values = []
    values_line = []
    
    for region in available_regions:
        commission = reviewsData.filter(region=region['region'])
        # print(commission)
        if commission.exists():
            commission = commission.aggregate(total_commission=Sum(Cast('commission', output_field=FloatField())))['total_commission']
        else:
            commission = 0
                # commission_trend.append({"date" : review_document['updated_at__date'].strftime('%d %b %Y'), "commission": float(gk['commission'].replace(',', '.'))})
        # region_commission_trend.append([region['region'], int(commission)])
        labels.append(region['region'])
        values.append(int(commission))
        values_line.append(int(int(commission)/int(total_commission)*100))
    region_commission_trend = {
        "labels" : labels,
        "values" : values,
        "values_line" : values_line
    }
    # top_regions = sorted(top_regions, key=lambda d: d['commission'], reverse=True)
    # Advisor wise Trend
    available_advisors = UserAccount.objects.filter(userType=6).values()
    top_advisors = {}
    labels = []
    emails = []
    id_numbers = []
    values = []
    for advisor in available_advisors:
        commission = reviewsData.filter(advisor=advisor['id'])
        if commission.exists():
            commission = commission.aggregate(total_commission=Sum(Cast('commission', output_field=FloatField())))['total_commission']
        else:
            commission = 0

                # commission_trend.append({"date" : review_document['updated_at__date'].strftime('%d %b %Y'), "commission": float(gk['commission'].replace(',', '.'))})
        if commission != 0:
            advisor_profile = user_profile.objects.filter(user=advisor['id'])
            name = f"{advisor['first_name']} {advisor['last_name']}"
            id = ""
            if advisor_profile.exists():
                name = advisor_profile.first().Full_Name
                name = name
                id = advisor_profile.first().ID_Number
            labels.append(name)
            emails.append(advisor['email'])
            id_numbers.append(id)
            values.append(int(commission))
    top_advisors = {
        "labels" : labels,
        "emails" : emails,
        "id_numbers" : id_numbers,
        "values" : values
    }
    # top_advisors = sorted(top_advisors, key=lambda d: d['commission'], reverse=True)
    # Business Type wise Trend
    businessType_commission_trend = {}
    labels = []
    values = []
    values_line = []
    business_total_commission = 0
    for i in range(1,16):
        commission = reviewsData.filter(businessType=i)
        if commission.exists():
            commission = commission.aggregate(total_commission=Sum(Cast('commission', output_field=FloatField())))['total_commission']
        else:
            commission = 0
        business_total_commission += int(commission)
        if i == 1:
            businessType = "Business Assurance"
        if i == 2:
            businessType = "Comm release"
        if i == 3:
            businessType = "Employee Benefits"
        if i == 4:
            businessType = "Funeral"
        if i == 5:
            businessType = "GAP Cover"
        if i == 6:
            businessType = "Recurring - Investment"
        if i == 7:
            businessType = "Lumpsum - Investment"
        if i == 8:
            businessType = "Investment- Both"
        if i == 9:
            businessType = "Medical Aid"
        if i == 10:
            businessType = "Other"
        if i == 11:
            businessType = "Will"
        if i == 12:
            businessType = "Risk"
        if i == 13:
            businessType = "ST Re-issued/instated"
        if i == 14:
            businessType = "Short Term Commercial"
        if i == 15:
            businessType = "Short Term Personal"    
        labels.append(businessType)
        values.append(int(commission))
        values_line.append(int(int(commission)/int(total_commission)*100))
    businessType_commission_trend = {
        "labels" : labels,
        "values" : values,
        "values_line" : values_line,
    }
        
    # plt.bar(data['labels'], data['values'], color='#5ABCFC')
    fig, ax1 = plt.subplots()
    ax1.plot(commission_trend['labels'], commission_trend['values'], color='#5ABCFC')
    ax1.set_xlabel('Date')
    ax1.set_ylabel('Total Commission', color='#5ABCFC')
    ax1.tick_params(axis='y', labelcolor='#5ABCFC')
    ax1.yaxis.set_major_formatter(currency_formatter)

    plt.xlabel('Date')
    plt.ylabel('Percentage of Commission')
    plt.title('Commission Sum Trend')

    # plt.bar(commission_trend['labels'], commission_trend['values'], color='#5ABCFC')
    plt.xticks(rotation=45, ha='right')  # Adjust the rotation angle and alignment as needed
    # Adjust subplot parameters to remove padding
    plt.subplots_adjust(left=0.1, right=0.9, top=0.9, bottom=0.2)  # Adjust as needed

    # Adjust figure size to increase height
    plt.gcf().set_size_inches(7, 6)  # Adjust the width and height as needed

    # Save the chart to a BytesIO object
    img_bytes = io.BytesIO()
    plt.savefig(img_bytes, format='png', bbox_inches='tight')
    img_bytes.seek(0)



    # Create a Word document
    doc = Document()

    # Add heading 1 for "SFP Operations Update Q4:2023"
    cover_page_heading = doc.add_heading("SFP Commission Report Export", level=1)
    cover_page_heading.style.font.name = 'Calibri'  # Change font to Calibri Light
    cover_page_heading.style.font.size = Pt(28) 
    cover_page_heading.style.font.color.rgb = RGBColor(0,0,0) 

    # Add the date
    doc.add_paragraph().add_run(f"Date: {datetime.now().strftime('%d %b %Y')}").bold = True

    # Add 5 line breaks
    for _ in range(5):
        doc.add_paragraph()

    # Add name
    # user = request.user
    requesting_user = user.first_name + " " + user.last_name
    user_profile_data = user_profile.objects.filter(user=user.pk)
    if user_profile_data.exists():
        requesting_user = user_profile_data.first().Full_Name
    doc.add_paragraph().add_run(requesting_user).bold = True
    # Add additional content to the cover page as needed

    # Add another page break to start content on the next page
    doc.add_page_break()

    heading = doc.add_heading('Compliance advice monitoring', level=1)
    heading.style.font.name = 'Calibri'  # Change font to Calibri Light
    heading.style.font.size = Pt(14) 
    heading.style.font.color.rgb = RGBColor(0,0,0) 
    heading.style.font.bold = False 

    table = doc.add_table(rows=1, cols=4)
    table.alignment = 1  # Center alignment
    cell = table.cell(0, 0)
    # cell.vertical_alignment = docx.enum.text.WD_ALIGN_VERTICAL.CENTER
    cell_paragraph = cell.paragraphs[0]
    cell_paragraph.alignment = 1  # Center alignment
    cell_paragraph.style.font.size = Pt(20)
    cell_paragraph.style.font.color.rgb = RGBColor(0, 0, 0)  # White color
    cell_paragraph.style.font.bold = True
    cell_paragraph.style.font.name = 'Arial'
    cell_paragraph.alignment = 1  # Center alignment
    run1 = cell_paragraph.add_run(babel.numbers.format_number(total_reviews,locale='en_ZA'))
    run1.font.size = Pt(15)
    run1 = cell_paragraph.add_run('\nTotal Reviews')
    run1.font.size = Pt(8)

    cell = table.cell(0, 1)
    # cell.vertical_alignment = docx.enum.text.WD_ALIGN_VERTICAL.CENTER
    cell_paragraph = cell.paragraphs[0]
    cell_paragraph.alignment = 1  # Center alignment
    cell_paragraph.style.font.size = Pt(20)
    cell_paragraph.style.font.color.rgb = RGBColor(0, 0, 0)  # White color
    cell_paragraph.style.font.bold = True
    cell_paragraph.style.font.name = 'Arial'
    cell_paragraph.alignment = 1  # Center alignment
    run1 = cell_paragraph.add_run(babel.numbers.format_currency(int(total_commission), 'ZAR', locale='en_ZA', currency_digits=False))
    run1.font.size = Pt(15)
    run1 = cell_paragraph.add_run('\nTotal Commission')
    run1.font.size = Pt(8)

    cell = table.cell(0, 2)
    # cell.vertical_alignment = docx.enum.text.WD_ALIGN_VERTICAL.CENTER
    cell_paragraph = cell.paragraphs[0]
    cell_paragraph.alignment = 1  # Center alignment
    cell_paragraph.style.font.size = Pt(20)
    cell_paragraph.style.font.color.rgb = RGBColor(0, 0, 0)  # White color
    cell_paragraph.style.font.bold = True
    cell_paragraph.style.font.name = 'Arial'
    cell_paragraph.alignment = 1  # Center alignment
    run1 = cell_paragraph.add_run(babel.numbers.format_number(total_regions,locale='en_ZA'))
    run1.font.size = Pt(15)
    run1 = cell_paragraph.add_run('\nRegions')
    run1.font.size = Pt(8)

    cell = table.cell(0, 3)
    # cell.vertical_alignment = docx.enum.text.WD_ALIGN_VERTICAL.CENTER
    cell_paragraph = cell.paragraphs[0]
    cell_paragraph.alignment = 1  # Center alignment
    cell_paragraph.style.font.size = Pt(20)
    cell_paragraph.style.font.color.rgb = RGBColor(0, 0, 0)  # White color
    cell_paragraph.style.font.bold = True
    cell_paragraph.style.font.name = 'Arial'
    cell_paragraph.alignment = 1  # Center alignment
    run1 = cell_paragraph.add_run(babel.numbers.format_number(total_advisors,locale='en_ZA'))
    run1.font.size = Pt(15)
    run1 = cell_paragraph.add_run('\nAdvisors')
    run1.font.size = Pt(8)


    heading = doc.add_heading('Commission Trend', level=1)
    heading.style.font.name = 'Calibri'  # Change font to Calibri Light
    heading.style.font.size = Pt(14) 
    heading.style.font.color.rgb = RGBColor(0,0,0) 
    heading.style.font.bold = False 


    # Add a table with headers
    commission_table = doc.add_table(rows=1, cols=2)
    commission_hdr_cells = commission_table.rows[0].cells
    commission_hdr_cells[0].text = 'Date'
    commission_hdr_cells[1].text = 'Total Commission'

    # Add commission_trend to the table
    for cell in commission_hdr_cells:
        cell.paragraphs[0].style.font.name = 'Calibri'
        cell.paragraphs[0].style.font.size = Pt(20)
        cell.paragraphs[0].style.font.bold = True
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Add commission_trend to the table
    for label, value in zip(commission_trend["labels"], commission_trend["values"]):
        row_cells = commission_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = babel.numbers.format_currency(int(value), 'ZAR', locale='en_ZA', currency_digits=False)
        
    # Change font to Calibri and set font size to 12 for commission_trend cells
    for row in commission_table.rows[1:]:
        for cell in row.cells:
            cell.paragraphs[0].style.font.name = 'Calibri'
            cell.paragraphs[0].style.font.size = Pt(12)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    image = doc.add_picture(img_bytes)
    doc.add_page_break()

    heading = doc.add_heading('Region Based Trend', level=1)
    heading.style.font.name = 'Calibri'  # Change font to Calibri Light
    heading.style.font.size = Pt(14) 
    heading.style.font.color.rgb = RGBColor(0,0,0) 
    heading.style.font.bold = False 



    # plt.bar(data['labels'], data['values'], color='#5ABCFC')
    fig, ax1 = plt.subplots()

    # Plot bar chart on primary y-axis
    ax1.bar(region_commission_trend['labels'], region_commission_trend['values'], color='#5ABCFC')
    ax1.set_xlabel('Regions')
    ax1.set_ylabel('Commission', color='#5ABCFC')
    ax1.tick_params(axis='y', labelcolor='#5ABCFC')

    # Rotate x-axis labels
    plt.xticks(rotation=45, ha='right')

    # Create a secondary y-axis and plot line chart
    ax2 = ax1.twinx()
    ax2.plot(region_commission_trend['labels'], region_commission_trend['values_line'], color='red', marker='o')
    ax2.set_ylabel('Percentage of Commission', color='red')
    ax2.tick_params(axis='y', labelcolor='red')

    ax1.yaxis.set_major_formatter(currency_formatter)
    ax2.yaxis.set_major_formatter(ScalarFormatter(useMathText=False))

    # Adjust subplot parameters to remove padding
    plt.subplots_adjust(left=0.1, right=0.9, top=0.9, bottom=0.2)

    # Adjust figure size to increase height
    plt.gcf().set_size_inches(7, 6)

    plt.title('Region Based Trend')

    # Save the chart to a BytesIO object
    region_img_bytes = io.BytesIO()
    plt.savefig(region_img_bytes, format='png', bbox_inches='tight')
    region_img_bytes.seek(0)

    # Add a table with headers
    bt_table = doc.add_table(rows=1, cols=2)
    bt_hdr_cells = bt_table.rows[0].cells
    bt_hdr_cells[0].text = 'Region'
    bt_hdr_cells[1].text = 'Total Commission'

    # Add data to the table
    for cell in bt_hdr_cells:
        cell.paragraphs[0].style.font.name = 'Calibri'
        cell.paragraphs[0].style.font.size = Pt(20)
        cell.paragraphs[0].style.font.bold = True
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Add data to the table
    for label, value in zip(region_commission_trend["labels"], region_commission_trend["values"]):
        row_cells = bt_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = babel.numbers.format_currency(int(value), 'ZAR', locale='en_ZA', currency_digits=False)
        
    # Change font to Calibri and set font size to 12 for data cells
    for row in bt_table.rows[1:]:
        for cell in row.cells:
            cell.paragraphs[0].style.font.name = 'Calibri'
            cell.paragraphs[0].style.font.size = Pt(12)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    image = doc.add_picture(region_img_bytes)
    doc.add_page_break()

    heading = doc.add_heading('Business Type Trend', level=1)
    heading.style.font.name = 'Calibri'  # Change font to Calibri Light
    heading.style.font.size = Pt(14) 
    heading.style.font.color.rgb = RGBColor(0,0,0) 
    heading.style.font.bold = False 
    # plt.bar(data['labels'], data['values'], color='#5ABCFC')
    fig, ax1 = plt.subplots()
    ax1.bar(businessType_commission_trend['labels'], businessType_commission_trend['values'], color='#5ABCFC')
    ax1.set_xlabel('Business Type')
    ax1.set_ylabel('Total Commission', color='#5ABCFC')
    ax1.tick_params(axis='y', labelcolor='#5ABCFC')

    # Rotate x-axis labels
    plt.xticks(rotation=45, ha='right')

    # Create a secondary y-axis and plot line chart
    ax2 = ax1.twinx()
    ax2.plot(businessType_commission_trend['labels'], businessType_commission_trend['values_line'], color='red', marker='o')
    ax2.set_ylabel('Percentage', color='red')
    ax2.tick_params(axis='y', labelcolor='red')

    ax1.yaxis.set_major_formatter(currency_formatter)
    ax2.yaxis.set_major_formatter(ScalarFormatter(useMathText=False))

    plt.xlabel('Date')
    plt.ylabel('Values')
    plt.title('Business Type Based Trend')

    # plt.bar(businessType_commission_trend['labels'], businessType_commission_trend['values'], color='#5ABCFC')
    plt.xticks(rotation=45, ha='right')  # Adjust the rotation angle and alignment as needed

    # Adjust subplot parameters to remove padding
    plt.subplots_adjust(left=0.1, right=0.9, top=0.9, bottom=0.2)  # Adjust as needed

    # Adjust figure size to increase height
    plt.gcf().set_size_inches(7, 6)  # Adjust the width and height as needed

    # Save the chart to a BytesIO object
    bt__img_bytes = io.BytesIO()
    plt.savefig(bt__img_bytes, format='png', bbox_inches='tight')
    bt__img_bytes.seek(0)

    # Add a table with headers
    region_table = doc.add_table(rows=1, cols=2)
    region_hdr_cells = region_table.rows[0].cells
    region_hdr_cells[0].text = 'Business Type'
    region_hdr_cells[1].text = 'Total Commission'

    # Add businessType_commission_trend to the table
    for cell in region_hdr_cells:
        cell.paragraphs[0].style.font.name = 'Calibri'
        cell.paragraphs[0].style.font.size = Pt(16)
        cell.paragraphs[0].style.font.bold = True
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Add businessType_commission_trend to the table
    for label, value in zip(businessType_commission_trend["labels"], businessType_commission_trend["values"]):
        row_cells = region_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = babel.numbers.format_currency(int(value), 'ZAR', locale='en_ZA', currency_digits=False)
        
    # Change font to Calibri and set font size to 12 for data cells
    for row in region_table.rows[1:]:
        for cell in row.cells:
            cell.paragraphs[0].style.font.name = 'Calibri'
            cell.paragraphs[0].style.font.size = Pt(10)
            cell.paragraphs[0].style.font.bold = False
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)


    image = doc.add_picture(bt__img_bytes)
    doc.add_page_break()

    heading = doc.add_heading('Advisor Data', level=1)
    heading.style.font.name = 'Calibri'  # Change font to Calibri Light
    heading.style.font.size = Pt(14) 
    heading.style.font.color.rgb = RGBColor(0,0,0) 
    heading.style.font.bold = False 
    

    # Add a table with headers
    advisor_table = doc.add_table(rows=1, cols=4)
    advisor_hdr_cells = advisor_table.rows[0].cells
    advisor_hdr_cells[0].text = 'Advisor'
    advisor_hdr_cells[1].text = 'Email'
    advisor_hdr_cells[2].text = 'ID Number'
    advisor_hdr_cells[3].text = 'Total Commission'

    # Add businessType_commission_trend to the table
    for cell in advisor_hdr_cells:
        cell.paragraphs[0].style.font.name = 'Calibri'
        cell.paragraphs[0].style.font.size = Pt(20)
        cell.paragraphs[0].style.font.bold = True
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Add businessType_commission_trend to the table
    for label, email, id, value in zip(top_advisors["labels"], top_advisors["emails"], top_advisors["id_numbers"], top_advisors["values"]):
        row_cells = advisor_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(email)
        row_cells[2].text = str(id)
        row_cells[3].text = babel.numbers.format_currency(int(value), 'ZAR', locale='en_ZA', currency_digits=False)
        
    # Change font to Calibri and set font size to 12 for data cells
    for row in advisor_table.rows[1:]:
        for cell in row.cells:
            cell.paragraphs[0].style.font.name = 'Calibri'
            cell.paragraphs[0].style.font.size = Pt(10)
            cell.paragraphs[0].style.font.bold = False
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)


    doc.add_page_break()


    # cell.fill.solid()
    # cell.fill.fore_color.rgb = RGBColor(32, 78, 120)  
    # Save the Word document to the static folder
    report_path = os.path.join(f'static/reports/{file_initial_name} Report.docx')
    doc.save(report_path)
    downloading_link = env('DJANGO_BACKEND_URL') + "/" + report_path
    notificationData = {
        "account" : user.pk,
        "notificationType" : 2,        
        "user" : user.pk,
        "title" : title,
        "message" : f"""
            <div style="max-width: 600px; margin: 20px auto; padding: 20px; background-color: #fff; border-radius: 8px; box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);">
                <h3 class='text-center' style="color: #333; margin-bottom: 20px;">{title}</h3>
                <div style="background-color: #FFCC00; color: #333; padding: 15px; border-radius: 4px; margin-bottom: 20px;">
                    <p style="margin: 0;">Dear {requesting_user},</p>
                    <p style="margin: 0;">Your requested commission insights export is ready to be downloaded.</p>
                    <div class='text-center'>
                        <a id="downloading_link" href="{downloading_link}" target="_blank" class="btn btn-primary btn-sm btn-sfp">Download</a>
                    </div>
                </div>
                <p style="color: #666; line-height: 1.6; margin-bottom: 15px;">If you have any questions or concerns feel free to reach out.</p>
                <p style="color: #666; line-height: 1.6; margin-bottom: 15px;">Best regards,<br>Succession</p>
            </div>
        """,
        "status" : False,
    }
    serializer = NotificationsSerializer(data=notificationData)
    if serializer.is_valid():
        serializer.create(serializer.validated_data)
    else:
        logger.info(serializer.errors)
    notificationData = {
        "account" : user.pk,
        "notificationType" : 6,        
        "user" : user.pk,
        "title" : title,
        "message" : f"{message} is ready.",
        "downloading_link" : report_path,
        "status" : False
    }
    serializer = NotificationsSerializer(data=notificationData)
    if serializer.is_valid():
        serializer.create(serializer.validated_data)
    else:
        logger.info(serializer.errors)

from django.template.loader import get_template
from django.core.mail import send_mail, EmailMultiAlternatives
from django.contrib.auth.tokens import default_token_generator
from djoser import utils

@shared_task
def sendResetPasswordEmail(email):
    user = UserAccount.objects.filter(email=email)
    if user.exists():
        from_email = 'SFP Support <support@sfponline.co.za>'
        user = user.first()
        data = {}
        data['protocol'] = 'https'
        data["site"] = "SFP Online"
        user_data = user_profile.objects.filter(user=user.pk)
        if user_data.exists():
            data["user_name"] = user_data.first().Full_Name
        else:
            data["user_name"] = user.first_name + " " + user.last_name
        data["URL"] = env('FRONTEND_URL') +"/auth/reset-password-confirm"
        data["uid"] = utils.encode_uid(user.pk)
        data["token"] = default_token_generator.make_token(user)
        template = get_template('password_reset.html').render(context=data)
        # print("Hi "+user[0]['email'])
        msg = EmailMultiAlternatives(
            'Password Reset Request',
            None, # This is the text context, just send None or Send a string message
            from_email,
            [email],
        )
        msg.attach_alternative(template, "text/html")
        msg.send(fail_silently=False)